from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup
from docx import Document

from src.features.rag.application.metadata import (
    detect_heading,
    detect_language,
    extract_hierarchy,
    normalize_text,
    parse_filename_metadata,
    strip_control_chars,
    strip_private_state,
    update_structure,
)
from src.features.rag.application.toc import TocIndex

SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".json", ".jsonl", ".html", ".htm", ".pdf", ".docx", ".csv", ".yaml", ".yml"
}


@dataclass(frozen=True)
class LoadedDocument:
    text: str
    normalized_text: str
    metadata: dict[str, Any]


def _render_structured_json(obj: Any, path: list[str] | None = None) -> list[LoadedDocument]:
    """Structured data is authoritative: only declared metadata keys are copied."""
    path = path or []
    docs: list[LoadedDocument] = []
    if isinstance(obj, dict):
        normalized = {str(k).casefold(): v for k, v in obj.items()}
        content_keys = [k for k in obj if str(k).casefold() in {"content", "text", "body", "description"}]
        if content_keys:
            content = "\n".join(str(obj[k]).strip() for k in content_keys if str(obj[k]).strip())
            if content:
                meta: dict[str, Any] = {}
                for key in ("subject", "grade", "chapter", "lesson", "section", "topic", "language"):
                    if key in normalized and not isinstance(normalized[key], (dict, list)):
                        meta[key] = normalized[key]
                header = "\n".join(path[-6:])
                text = f"{header}\n{content}".strip() if header else content
                docs.append(LoadedDocument(text=text, normalized_text=normalize_text(text), metadata=meta))
        for key, value in obj.items():
            if key in content_keys:
                continue
            next_path = path + [str(key)]
            if isinstance(value, dict):
                docs.extend(_render_structured_json(value, next_path))
            elif isinstance(value, list):
                for idx, item in enumerate(value, 1):
                    docs.extend(_render_structured_json(item, next_path + [f"Item {idx}"]))
    elif isinstance(obj, list):
        for idx, item in enumerate(obj, 1):
            docs.extend(_render_structured_json(item, path + [f"Item {idx}"]))
    return docs


def _apply_page_headings(lines: list[str], state: dict[str, Any]) -> tuple[list[str], dict[str, Any]]:
    output: list[str] = []
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            if output and output[-1] != "":
                output.append("")
            continue
        markdown = detect_heading(cleaned)
        update_structure(state, cleaned)
        output.append(cleaned)
    while output and output[-1] == "":
        output.pop()
    return output, state


def _load_pdf_with_llamaparse(path: Path, base: dict[str, Any], seed: dict[str, Any]) -> list[LoadedDocument] | None:
    from src.features.rag.infrastructure.config import get_settings

    settings = get_settings()
    if not settings.llama_parse_enabled or not settings.llama_cloud_api_key:
        return None

    # LlamaParse migrated to a tier-based v2 API and a new SDK package (`llama_cloud`,
    # replacing the deprecated `llama_cloud_services`). The old parse_mode= kwarg (and
    # the vendor-multimodal-model options that rode along with it) is v1-only and now
    # returns HTTP 410 Gone: "This parsing mode is no longer supported. Use tiers
    # instead — set `tier` to one of `cost_effective`, `agentic`, or `agentic_plus`."
    #
    # We standardize on `agentic`, not the cheaper `cost_effective`: this curriculum's
    # PDFs carry a known embedded-font ToUnicode bug that garbles Arabic ligatures
    # (lam+alef-family) in the text layer itself (see
    # detect_arabic_ligature_corruption() in document_metadata.py). `agentic` is the
    # tier that actually runs OCR — reading the page as an image — which bypasses that
    # broken text layer entirely; `cost_effective` is not documented to. These books
    # are also table/exercise-layout heavy (multi-column, boxed answer choices), which
    # `agentic` handles explicitly.
    #
    # Cost Optimizer (processing_options.cost_optimizer.enable) is layered on top of
    # `agentic`/`agentic_plus` specifically — it is NOT a standalone tier and cannot be
    # combined with `cost_effective`. It pre-analyzes each page's complexity and
    # downgrades simple prose pages to cost_effective pricing automatically, so most of
    # a text-heavy book still bills cheap while tables/complex layout get full agentic
    # treatment — without you having to guess up front which pages need which tier.
    #
    # NOTE: this rewrite is based on LlamaIndex's current (2026) v2 docs, since this
    # sandbox has no network access to pip-install `llama-cloud` and exercise the
    # actual client — attribute names below (e.g. `result.markdown.pages[i].markdown`)
    # are inferred from documented Go/Java examples with matching field names and are
    # NOT verified against the installed Python SDK. If this throws an AttributeError,
    # run `python -c "from llama_cloud import LlamaCloud; help(LlamaCloud)"` (and
    # `help(client.parsing.parse)`) after installing, and adjust the field access below
    # to match — the polling/upload flow and processing_options shape should hold even
    # if an exact attribute name differs.
    import time
    import requests

    headers = {"Authorization": f"Bearer {settings.llama_cloud_api_key}"}
    with open(str(path), "rb") as f:
        files = {"file": (path.name, f, "application/pdf")}
        data = {
            "tier": "agentic",
            "version": "latest",
            "language": "ar",
        }
        resp = requests.post(
            "https://api.cloud.llamaindex.ai/api/parsing/upload",
            headers=headers,
            files=files,
            data=data,
            timeout=120,
        )
    if resp.status_code != 200:
        raise RuntimeError(f"LlamaParse upload failed ({resp.status_code}): {resp.text}")

    job_id = resp.json().get("id")
    if not job_id:
        raise RuntimeError(f"LlamaParse returned no job ID: {resp.text}")

    max_wait = 300
    start_time = time.time()
    while time.time() - start_time < max_wait:
        status_resp = requests.get(f"https://api.cloud.llamaindex.ai/api/parsing/job/{job_id}", headers=headers, timeout=30)
        status_data = status_resp.json() if status_resp.status_code == 200 else {}
        status_str = str(status_data.get("status", "")).upper()
        if status_str in ("SUCCESS", "COMPLETED"):
            break
        if status_str in ("ERROR", "FAILED"):
            raise RuntimeError(f"LlamaParse job {job_id} failed: {status_data.get('error_message', status_str)}")
        time.sleep(2)

    result_resp = requests.get(f"https://api.cloud.llamaindex.ai/api/parsing/job/{job_id}/result/json", headers=headers, timeout=60)
    if result_resp.status_code == 200:
        result_json = result_resp.json()
        pages = result_json.get("pages", []) or [{"page": 1, "text": result_json.get("markdown", "")}]
    else:
        result_resp = requests.get(f"https://api.cloud.llamaindex.ai/api/parsing/job/{job_id}/result/markdown", headers=headers, timeout=60)
        markdown_text = result_resp.json().get("markdown", "") if result_resp.status_code == 200 else ""
        pages = [{"page": 1, "text": markdown_text}]

    if not pages:
        return None

    page_texts = [str(p.get("text") or p.get("markdown") or "") for p in pages]
    if not any(t.strip() for t in page_texts):
        return None

    toc_index = TocIndex.build(page_texts)

    state = dict(seed)
    docs: list[LoadedDocument] = []
    for idx, (page_obj, text) in enumerate(zip(pages, page_texts), 1):
        text = text.strip()
        if not text:
            continue
        page_no = getattr(page_obj, "page", None) or getattr(page_obj, "index", None) or idx
        try:
            page_no = int(page_no)
        except (TypeError, ValueError):
            page_no = idx

        # LlamaParse commonly returns markdown headings. Keep a generic heading state,
        # without deciding that a numbered heading is automatically a chapter or lesson.
        page_state = dict(state)
        for raw_line in text.splitlines():
            line = strip_control_chars(raw_line.strip())
            if not line:
                continue
            update_structure(state, line, toc_index=toc_index)

        docs.append(
            LoadedDocument(
                text=text,
                normalized_text=normalize_text(text),
                metadata={
                    **base,
                    **strip_private_state(page_state),
                    "page": page_no,
                    "parser": "llamaparse",
                    "language": detect_language(text),
                },
            )
        )
    return docs


def _order_page_lines(
    page_lines: list[tuple[str, float, bool, float, float]],
    page_width: float,
) -> list[tuple[str, float, bool, float, float]]:
    """Reconstruct reading order from PyMuPDF line boxes.

    A flat sort by (y, x) — the previous behaviour — assumes the whole page reads
    left-to-right, top-to-bottom as one stream. That's wrong for genuine multi-column
    pages: two columns sharing the same row band get interleaved line-by-line (col1 row1,
    col2 row1, col1 row2, col2 row2, ...) instead of reading column 1 fully before column 2.

    This clusters lines into left-to-right column bands by a gap in their x0 position, but
    only commits to a multi-column split when the gap is wide relative to the page and each
    side has enough lines to be a real column (not a stray page-number/logo box) — otherwise
    it falls back to the previous flat (y, x) sort, which stays correct for the common case
    of a single-column page or a cover page with a couple of stray boxes.
    """
    if not page_lines:
        return []
    if len(page_lines) < 6 or page_width <= 0:
        return sorted(page_lines, key=lambda x: (round(x[3], 1), round(x[4], 1)))

    xs = sorted(x0 for *_, x0 in page_lines)
    best_gap, best_split = 0.0, None
    for i in range(1, len(xs)):
        gap = xs[i] - xs[i - 1]
        if gap > best_gap:
            best_gap, best_split = gap, (xs[i - 1] + xs[i]) / 2

    min_column_lines = max(3, len(page_lines) // 6)
    if best_split is not None and best_gap >= page_width * 0.08:
        left = [ln for ln in page_lines if ln[4] < best_split]
        right = [ln for ln in page_lines if ln[4] >= best_split]
        if len(left) >= min_column_lines and len(right) >= min_column_lines:
            left.sort(key=lambda x: (round(x[3], 1), round(x[4], 1)))
            right.sort(key=lambda x: (round(x[3], 1), round(x[4], 1)))
            return left + right

    return sorted(page_lines, key=lambda x: (round(x[3], 1), round(x[4], 1)))


def _load_pdf_with_pymupdf(path: Path, base: dict[str, Any], seed: dict[str, Any]) -> list[LoadedDocument]:
    try:
        import fitz
    except ImportError:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        page_texts = [page.extract_text() or "" for page in reader.pages]
        # See the LlamaParse loader above for why this is built up front, from every
        # page's raw text, before the heading-tracking walk begins.
        toc_index = TocIndex.build(page_texts)
        state = dict(seed)
        docs: list[LoadedDocument] = []
        for page_no, text in enumerate(page_texts, 1):
            page_state = dict(state)
            for line in text.splitlines():
                cleaned = strip_control_chars(line.strip())
                if cleaned:
                    update_structure(state, cleaned, toc_index=toc_index)
            docs.append(
                LoadedDocument(
                    text=text,
                    normalized_text=normalize_text(text),
                    metadata={**base, **strip_private_state(page_state), "page": page_no, "parser": "pypdf", "language": detect_language(text)},
                )
            )
        return docs

    pdf = fitz.open(str(path))
    state = dict(seed)
    docs: list[LoadedDocument] = []
    try:
        page_texts = [p.get_text("text") or "" for p in pdf]
        toc_index = TocIndex.build(page_texts)
        for page_no, page in enumerate(pdf, 1):
            page_state = dict(state)
            page_lines: list[tuple[str, float, bool, float, float]] = []
            data = page.get_text("dict")
            for block in data.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = [s for s in line.get("spans", []) if str(s.get("text", "")).strip()]
                    if not spans:
                        continue
                    txt = strip_control_chars("".join(str(s.get("text", "")) for s in spans).strip())
                    if not txt:
                        continue
                    sizes = [float(s.get("size", 0) or 0) for s in spans]
                    font_size = sum(sizes) / max(len(sizes), 1)
                    bold = any("bold" in str(s.get("font", "")).casefold() or int(s.get("flags", 0) or 0) & 16 for s in spans)
                    bbox = line.get("bbox", [0, 0, 0, 0])
                    page_lines.append((txt, font_size, bold, float(bbox[1]), float(bbox[0])))
            sizes = sorted(x[1] for x in page_lines if x[1] > 0)
            median = sizes[len(sizes) // 2] if sizes else 0
            page_lines = _order_page_lines(page_lines, float(page.rect.width))
            output_lines: list[str] = []
            for txt, size, bold, _, _ in page_lines:
                # Style signals are only used to decide whether a line is a heading; semantic level
                # comes from generic numbering/markdown/explicit labels.
                candidate = detect_heading(txt)
                style_heading = bool(candidate and (bold or (median and size >= median * 1.18)))
                explicit = update_structure(state, txt, toc_index=toc_index)
                if explicit or style_heading:
                    update_structure(state, txt, toc_index=toc_index)
                output_lines.append(txt)
            page_text = "\n".join(output_lines).strip()
            docs.append(
                LoadedDocument(
                    text=page_text,
                    normalized_text=normalize_text(page_text),
                    metadata={**base, **strip_private_state(page_state), "page": page_no, "parser": "pymupdf", "language": detect_language(page_text)},
                )
            )
    finally:
        pdf.close()
    return docs


def _pymupdf_cover_text(path: Path, max_pages: int = 8) -> str | None:
    """Best-effort plain-text extraction of just the first few pages via PyMuPDF, used only
    to feed document-level metadata extraction (title/subject/grade) — never for chunk
    content. Cover/title pages are short, position-and-font-structured text that PyMuPDF's
    line-box extraction handles well; LlamaParse's markdown reflow can restructure that
    layout in ways that break the metadata regexes even when it's the better choice for the
    book's actual (often complex/tabular) body content.

    Always returns None instead of raising — this is a supplementary signal only, so any
    failure (fitz not installed, encrypted/corrupt PDF, etc.) must fall back silently to
    whatever text the primary loader already produced.
    """
    try:
        import fitz
    except ImportError:
        return None
    try:
        pdf = fitz.open(str(path))
    except Exception:
        return None
    try:
        pages_text: list[str] = []
        for page in list(pdf)[:max_pages]:
            page_lines: list[tuple[str, float, bool, float, float]] = []
            data = page.get_text("dict")
            for block in data.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = [s for s in line.get("spans", []) if str(s.get("text", "")).strip()]
                    if not spans:
                        continue
                    txt = "".join(str(s.get("text", "")) for s in spans).strip()
                    bbox = line.get("bbox", [0, 0, 0, 0])
                    page_lines.append((txt, 0.0, False, float(bbox[1]), float(bbox[0])))
            ordered = _order_page_lines(page_lines, float(page.rect.width))
            pages_text.append("\n".join(txt for txt, *_ in ordered))
        text = "\n".join(pages_text).strip()
        return text or None
    except Exception:
        return None
    finally:
        pdf.close()


def _load_pdf(path: Path, base: dict[str, Any], seed: dict[str, Any]) -> list[LoadedDocument]:
    try:
        parsed = _load_pdf_with_llamaparse(path, base, seed)
        if parsed:
            # LlamaParse handled the body; still prefer PyMuPDF's cleaner cover-page text
            # for document-level metadata extraction (see _pymupdf_cover_text). Stashed as
            # a private, non-metadata key on the first doc — the ingestion pipeline reads
            # and discards it before any of it reaches chunk metadata.
            cover_text = _pymupdf_cover_text(path)
            if cover_text and parsed:
                parsed[0] = LoadedDocument(
                    text=parsed[0].text,
                    normalized_text=parsed[0].normalized_text,
                    metadata={**parsed[0].metadata, "_metadata_source_text": cover_text},
                )
            return parsed
    except Exception as exc:
        from src.features.rag.utils.logger import get_logger
        get_logger(__name__).warning("LlamaParse failed for %s; falling back to PyMuPDF: %s", path.name, exc)
    return _load_pdf_with_pymupdf(path, base, seed)


def _load_docx(path: Path, base: dict[str, Any], seed: dict[str, Any]) -> LoadedDocument:
    doc = Document(str(path))
    state = dict(seed)
    lines: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").casefold() if p.style else ""
        if "heading" in style:
            level_match = next((int(x) for x in re.findall(r"heading\s*(\d+)", style)), 1) if "heading" in style else 1
            state["heading"] = text
            state["heading_level"] = level_match
            path_state = [p for p in (state.get("heading_path") or []) if p]
            state["heading_path"] = " > ".join(path_state[: level_match - 1] + [text])
        update_structure(state, text)
        lines.append(text)
    text = "\n".join(lines)
    return LoadedDocument(text=text, normalized_text=normalize_text(text), metadata={**base, **strip_private_state(state)})


def load_file(path: Path, *, extra_metadata: dict[str, Any] | None = None, file_reference_id: str | None = None) -> list[LoadedDocument]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")

    base = {"source": path.name, "source_type": suffix.lstrip("."), **parse_filename_metadata(path.name)}
    if extra_metadata:
        for key, value in extra_metadata.items():
            if value is not None:
                base[key] = value
    if file_reference_id:
        base["file_reference_id"] = file_reference_id

    seed = extract_hierarchy("", base)

    if suffix in {".json", ".yaml", ".yml"}:
        if suffix == ".json":
            obj = json.loads(path.read_text(encoding="utf-8"))
        else:
            import yaml
            obj = yaml.safe_load(path.read_text(encoding="utf-8"))
        docs = _render_structured_json(obj)
        if docs:
            return [LoadedDocument(d.text, d.normalized_text, {**base, **d.metadata}) for d in docs]
        text = json.dumps(obj, ensure_ascii=False, indent=2)
        return [LoadedDocument(text, normalize_text(text), base)]

    if suffix == ".jsonl":
        docs: list[LoadedDocument] = []
        for line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
            if not line.strip():
                continue
            obj = json.loads(line)
            docs.extend(_render_structured_json(obj))
        return [LoadedDocument(d.text, d.normalized_text, {**base, **d.metadata}) for d in docs]

    if suffix in {".txt", ".md", ".markdown"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [LoadedDocument(text, normalize_text(text), {**base, **extract_hierarchy(text, base)})]

    if suffix in {".html", ".htm"}:
        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        state = dict(seed)
        lines: list[str] = []
        for node in soup.find_all(["h1", "h2", "h3", "h4", "p", "li"]):
            text = node.get_text(" ", strip=True)
            if not text:
                continue
            if node.name.startswith("h"):
                level = int(node.name[1:])
                state["heading"] = text
                state["heading_level"] = level
                state["heading_path"] = " > ".join([p for p in str(state.get("heading_path") or "").split(" > ") if p][: level - 1] + [text])
            update_structure(state, text)
            lines.append(text)
        text = "\n".join(lines)
        return [LoadedDocument(text, normalize_text(text), {**base, **strip_private_state(state)})]

    if suffix == ".pdf":
        return _load_pdf(path, base, seed)
    if suffix == ".docx":
        return [_load_docx(path, base, seed)]
    if suffix == ".csv":
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
            rows = list(csv.reader(f))
        text = "\n".join(" | ".join(row) for row in rows)
        return [LoadedDocument(text, normalize_text(text), base)]

    raise ValueError(f"Unhandled suffix: {suffix}")


def discover_files(source: Path) -> list[Path]:
    if source.is_file():
        return [source]
    return sorted(p for p in source.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS)
